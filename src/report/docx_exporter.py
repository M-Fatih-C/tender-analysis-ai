"""
TenderAI Word/DOCX Rapor Üretici v1.0.

Analiz sonuçlarını profesyonel Word belgesi olarak dışa aktarır.
"""

import io
import logging
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Renk paleti
_COLORS = {
    "primary": RGBColor(0x66, 0x7E, 0xEA),
    "success": RGBColor(0x27, 0xAE, 0x60),
    "warning": RGBColor(0xF3, 0x9C, 0x12),
    "danger": RGBColor(0xE7, 0x4C, 0x3C),
    "dark": RGBColor(0x2C, 0x3E, 0x50),
    "muted": RGBColor(0x7F, 0x8C, 0x8D),
}


def generate_docx_report(
    analysis_result: dict,
    file_name: str = "İhale Analizi",
    company_name: str = "",
) -> bytes:
    """
    Analiz sonucunu Word/DOCX formatında döndür.

    Returns:
        bytes — .docx dosya içeriği
    """
    doc = Document()

    # Stil ayarları
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = _COLORS["dark"]

    # ── Kapak sayfası ──
    _add_cover_page(doc, file_name, company_name)

    # ── Yönetici Özeti ──
    _add_executive_summary(doc, analysis_result)

    # ── Risk Analizi ──
    _add_risk_analysis(doc, analysis_result)

    # ── Gerekli Belgeler ──
    _add_required_documents(doc, analysis_result)

    # ── Ceza Maddeleri ──
    _add_penalty_clauses(doc, analysis_result)

    # ── Mali Özet ──
    _add_financial_summary(doc, analysis_result)

    # ── Süre Analizi ──
    _add_timeline_analysis(doc, analysis_result)

    # ── Alt Bilgi ──
    _add_footer(doc)

    # Bytes çıktı
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==============================================================
# SECTIONS
# ==============================================================

def _add_cover_page(doc: Document, file_name: str, company_name: str) -> None:
    """Kapak sayfası."""
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TenderAI")
    run.font.size = Pt(36)
    run.font.color.rgb = _COLORS["primary"]
    run.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("İHALE ANALİZ RAPORU")
    run.font.size = Pt(16)
    run.font.color.rgb = _COLORS["muted"]

    doc.add_paragraph("")

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"📄 {file_name}")
    run.font.size = Pt(12)
    run.font.color.rgb = _COLORS["dark"]

    if company_name:
        co = doc.add_paragraph()
        co.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = co.add_run(f"🏢 {company_name}")
        run.font.size = Pt(11)
        run.font.color.rgb = _COLORS["muted"]

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"📅 {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = _COLORS["muted"]

    doc.add_page_break()


def _add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    """Bölüm başlık."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = _COLORS["primary"]


def _add_executive_summary(doc: Document, result: dict) -> None:
    """Yönetici özeti."""
    _add_section_heading(doc, "1. Yönetici Özeti")

    es = result.get("executive_summary", {})
    if isinstance(es, dict):
        ozet = es.get("ozet", "") or es.get("summary", "")
        tavsiye = es.get("tavsiye", "") or es.get("recommendation", "")
    else:
        ozet = str(es)[:300]
        tavsiye = ""

    score = result.get("risk_score", 0)
    level = result.get("risk_level", "—")

    # Risk skoru tablosu
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = f"Risk Skoru: {score}/100"
    hdr[1].text = f"Seviye: {level}"
    hdr[2].text = f"Model: {result.get('model_used', '—')}"
    _style_table_header(table)

    doc.add_paragraph("")

    if ozet:
        doc.add_paragraph(ozet)
    if tavsiye:
        p = doc.add_paragraph()
        run = p.add_run(f"💡 Tavsiye: {tavsiye}")
        run.font.italic = True
        run.font.color.rgb = _COLORS["primary"]


def _add_risk_analysis(doc: Document, result: dict) -> None:
    """Risk analizi bölümü."""
    _add_section_heading(doc, "2. Risk Analizi")

    ra = result.get("risk_analysis", {})
    if not isinstance(ra, dict):
        doc.add_paragraph("Risk analizi verisi mevcut değil.")
        return

    riskler = ra.get("riskler", [])
    if not riskler:
        doc.add_paragraph("Tespit edilen risk bulunmamaktadır.")
        return

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "Seviye"
    hdr[1].text = "Kategori"
    hdr[2].text = "Açıklama"
    hdr[3].text = "Madde"
    _style_table_header(table)

    for risk in riskler[:15]:
        if isinstance(risk, dict):
            row = table.add_row().cells
            row[0].text = risk.get("seviye", "—")
            row[1].text = risk.get("kategori", "—")
            row[2].text = risk.get("aciklama", "—")
            row[3].text = risk.get("madde_no", "—")


def _add_required_documents(doc: Document, result: dict) -> None:
    """Gerekli belgeler."""
    _add_section_heading(doc, "3. Gerekli Belgeler")

    docs_data = result.get("required_documents", {})
    if not isinstance(docs_data, dict):
        doc.add_paragraph("Belge listesi mevcut değil.")
        return

    zorunlu = docs_data.get("zorunlu_belgeler", [])
    if zorunlu:
        doc.add_heading("Zorunlu Belgeler", level=3)
        for d in zorunlu[:20]:
            if isinstance(d, dict):
                name = d.get("belge_adi", "—")
                desc = d.get("aciklama", "")
                doc.add_paragraph(f"✅ {name}" + (f" — {desc}" if desc else ""), style="List Bullet")


def _add_penalty_clauses(doc: Document, result: dict) -> None:
    """Ceza maddeleri."""
    _add_section_heading(doc, "4. Ceza Maddeleri")

    pc = result.get("penalty_clauses", {})
    if not isinstance(pc, dict):
        doc.add_paragraph("Ceza maddesi verisi mevcut değil.")
        return

    cezalar = pc.get("cezalar", [])
    if cezalar:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Ceza Türü"
        hdr[1].text = "Oran/Tutar"
        hdr[2].text = "Koşul"
        _style_table_header(table)

        for c in cezalar[:10]:
            if isinstance(c, dict):
                row = table.add_row().cells
                row[0].text = c.get("tur", "—")
                row[1].text = c.get("oran", "—")
                row[2].text = c.get("kosul", "—")


def _add_financial_summary(doc: Document, result: dict) -> None:
    """Mali özet."""
    _add_section_heading(doc, "5. Mali Özet")

    fin = result.get("financial_summary", {})
    if not isinstance(fin, dict):
        doc.add_paragraph("Mali analiz verisi mevcut değil.")
        return

    for key, label in [
        ("tahmini_bedel", "📊 Tahmini Bedel"),
        ("teminat_tutari", "💰 Teminat Tutarı"),
        ("avans_durumu", "💵 Avans Durumu"),
    ]:
        val = fin.get(key, "—")
        if val:
            doc.add_paragraph(f"{label}: {val}")


def _add_timeline_analysis(doc: Document, result: dict) -> None:
    """Süre analizi."""
    _add_section_heading(doc, "6. Süre Analizi")

    ta = result.get("timeline_analysis", {})
    if not isinstance(ta, dict):
        doc.add_paragraph("Süre analizi verisi mevcut değil.")
        return

    for key, label in [
        ("toplam_sure", "⏱️ Toplam Süre"),
        ("is_baslangic", "📅 İş Başlangıcı"),
        ("teslim_tarihi", "📅 Teslim Tarihi"),
        ("yer_teslim_suresi", "📅 Yer Teslim Süresi"),
    ]:
        val = ta.get(key, "")
        if val:
            doc.add_paragraph(f"{label}: {val}")


def _add_footer(doc: Document) -> None:
    """Alt bilgi."""
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("—— TenderAI v2.0 — Yapay Zeka İhale Analiz Platformu ——")
    run.font.size = Pt(8)
    run.font.color.rgb = _COLORS["muted"]
    run.font.italic = True


def _style_table_header(table) -> None:
    """Tablo başlığını renklendir."""
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Arka plan
        shading = cell._tc.get_or_add_tcPr()
        bg = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "667EEA",
        })
        shading.append(bg)
